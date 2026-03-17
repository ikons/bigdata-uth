#!/usr/bin/env python3
from __future__ import annotations

import argparse
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ACCENT_COLOR = RGBColor(0x1F, 0x4E, 0x79)
HEADER_BORDER_COLOR = "D9E2F2"
CODE_BACKGROUND = "EDEDED"
TOP_MARGIN_CM = 1.5
BOTTOM_MARGIN_CM = 1.5
LEFT_MARGIN_CM = 1.5
RIGHT_MARGIN_CM = 1.5
HEADER_DISTANCE_CM = 0.7
FOOTER_DISTANCE_CM = 0.7


def resolve_output_path(value: str) -> Path:
    raw_path = Path(value)
    if raw_path.is_absolute():
        return raw_path.resolve()
    return (Path.cwd() / raw_path).resolve()


def default_header_text() -> str:
    return Path(__file__).resolve().parent.parent.name


def get_or_add(parent, tag: str):
    element = parent.find(qn(tag))
    if element is None:
        element = OxmlElement(tag)
        parent.append(element)
    return element


def clear_body(document: Document) -> None:
    body = document._body._element
    sect_pr = deepcopy(document.sections[0]._sectPr)
    for child in list(body):
        body.remove(child)
    body.append(OxmlElement("w:p"))
    body.append(sect_pr)


def add_field(run, instruction: str, default_text: str = "") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    if default_text:
        text = OxmlElement("w:t")
        text.text = default_text
        run._r.append(text)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def add_preserved_instr_text(run, instruction: str) -> None:
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)


def add_simple_field(run, instruction: str, default_text: str = "") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)
    add_preserved_instr_text(run, instruction)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)
    if default_text:
        text = OxmlElement("w:t")
        text.text = default_text
        run._r.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def set_paragraph_border_bottom(paragraph, color: str) -> None:
    p_pr = get_or_add(paragraph._p, "w:pPr")
    borders = get_or_add(p_pr, "w:pBdr")
    bottom = get_or_add(borders, "w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def set_style_shading(style, fill: str, paragraph_style: bool) -> None:
    style_element = style._element
    target = get_or_add(style_element, "w:pPr" if paragraph_style else "w:rPr")
    shading = get_or_add(target, "w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)


def set_style_no_proof(style) -> None:
    r_pr = get_or_add(style._element, "w:rPr")
    get_or_add(r_pr, "w:noProof")


def set_update_fields(document: Document) -> None:
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def ensure_style(document: Document, name: str, style_type):
    try:
        return document.styles[name]
    except KeyError:
        return document.styles.add_style(name, style_type)


def build_template(output_path: Path, header_text: str) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Cm(TOP_MARGIN_CM)
    section.bottom_margin = Cm(BOTTOM_MARGIN_CM)
    section.left_margin = Cm(LEFT_MARGIN_CM)
    section.right_margin = Cm(RIGHT_MARGIN_CM)
    section.header_distance = Cm(HEADER_DISTANCE_CM)
    section.footer_distance = Cm(FOOTER_DISTANCE_CM)
    section.different_first_page_header_footer = True
    usable_width = section.page_width - section.left_margin - section.right_margin

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    title = document.styles["Title"]
    title.base_style = normal
    title.font.name = "Cambria"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = ACCENT_COLOR
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(18)

    heading1 = document.styles["Heading 1"]
    heading1.base_style = normal
    heading1.font.name = "Cambria"
    heading1.font.size = Pt(15)
    heading1.font.bold = True
    heading1.font.color.rgb = ACCENT_COLOR
    heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading1.paragraph_format.space_before = Pt(16)
    heading1.paragraph_format.space_after = Pt(6)

    heading2 = document.styles["Heading 2"]
    heading2.base_style = normal
    heading2.font.name = "Cambria"
    heading2.font.size = Pt(12.5)
    heading2.font.bold = True
    heading2.font.color.rgb = ACCENT_COLOR
    heading2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading2.paragraph_format.space_before = Pt(12)
    heading2.paragraph_format.space_after = Pt(4)

    heading3 = document.styles["Heading 3"]
    heading3.base_style = normal
    heading3.font.name = "Cambria"
    heading3.font.size = Pt(11)
    heading3.font.bold = True
    heading3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading3.paragraph_format.space_before = Pt(10)
    heading3.paragraph_format.space_after = Pt(3)

    contents_heading = ensure_style(document, "Contents Heading", WD_STYLE_TYPE.PARAGRAPH)
    contents_heading.base_style = normal
    contents_heading.font.name = "Cambria"
    contents_heading.font.size = Pt(14)
    contents_heading.font.bold = True
    contents_heading.font.color.rgb = ACCENT_COLOR
    contents_heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    contents_heading.paragraph_format.space_before = Pt(0)
    contents_heading.paragraph_format.space_after = Pt(8)

    for level, indent_cm in ((1, 0.0), (2, 0.5), (3, 1.0)):
        toc_style = ensure_style(document, f"Contents {level}", WD_STYLE_TYPE.PARAGRAPH)
        toc_style.base_style = normal
        toc_style.font.name = "Calibri"
        toc_style.font.size = Pt(10.5)
        toc_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        toc_style.paragraph_format.left_indent = Cm(indent_cm)
        toc_style.paragraph_format.space_before = Pt(0)
        toc_style.paragraph_format.space_after = Pt(2)
        toc_style.paragraph_format.line_spacing = 1.0

    source_code = ensure_style(document, "Source Code", WD_STYLE_TYPE.PARAGRAPH)
    source_code.base_style = normal
    source_code.style_id = "SourceCode"
    source_code.font.name = "Courier New"
    source_code.font.size = Pt(10)
    source_code.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    source_code.paragraph_format.space_before = Pt(0)
    source_code.paragraph_format.space_after = Pt(0)
    source_code.paragraph_format.line_spacing = 1.0
    set_style_shading(source_code, CODE_BACKGROUND, paragraph_style=True)

    verbatim_char = ensure_style(document, "Verbatim Char", WD_STYLE_TYPE.CHARACTER)
    verbatim_char.base_style = document.styles["Default Paragraph Font"]
    verbatim_char.style_id = "VerbatimChar"
    verbatim_char.font.name = "Courier New"
    verbatim_char.font.size = Pt(10)
    set_style_shading(verbatim_char, CODE_BACKGROUND, paragraph_style=False)
    set_style_no_proof(verbatim_char)

    for header in (section.header, section.first_page_header):
        header_paragraph = header.paragraphs[0]
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_paragraph.paragraph_format.tab_stops.add_tab_stop(
            usable_width,
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.SPACES,
        )
        header_run = header_paragraph.add_run(header_text)
        header_run.font.name = "Calibri"
        header_run.font.size = Pt(9)
        header_run.font.bold = True
        header_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        header_field_run = header_paragraph.add_run("\t")
        header_field_run.font.name = "Calibri"
        header_field_run.font.size = Pt(9)
        header_field_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        add_simple_field(header_field_run, ' STYLEREF "Heading 1" ', "Section")
        set_paragraph_border_bottom(header_paragraph, HEADER_BORDER_COLOR)

    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_paragraph.add_run()
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(9)
    add_field(footer_run, "PAGE", "1")

    first_page_footer_paragraph = section.first_page_footer.paragraphs[0]
    first_page_footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    clear_body(document)
    set_update_fields(document)
    document.save(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Create the DOCX reference template from scratch.")
    parser.add_argument(
        "--output",
        default="templates/reference.docx",
        help="Path to the generated reference template.",
    )
    parser.add_argument(
        "--header-text",
        default=default_header_text(),
        help="Header label used in the generated template.",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    output_path = resolve_output_path(args.output)
    build_template(output_path, args.header_text)
    print(f"Reference template created at {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
